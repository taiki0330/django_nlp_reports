from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('/Users/matsuzakidaiki/anaconda3/envs/django_nlp_env/reportsproject/reportsapp/docx_templates/ninsou.docx')
print(type(doc))

sentence = ['撮影者','本職']

for i in range(0, len(sentence)):
    doc.add_paragraph(sentence[i])

doc.add_heading('報告書')
doc.add_heading('報告書です', level=2)
# doc.add_page_break()

doc.save('ninsou2.docx')


#! new.docx -----------------------------------------------------------------
doc1 = Document()

create_date = doc1.add_paragraph('令和６年３月３日')
police_station = doc1.add_paragraph('福岡県東警察署')
police_l_1 = doc1.add_paragraph('司法警察員　警視正')
police_l_2 = doc1.add_paragraph('東　太郎　殿')
title = doc1.add_paragraph('写真撮影報告書')
title_content = doc1.add_paragraph('令和６年３月１日、当署館内において発生した被疑者天神太郎にかかる暴行被疑事件につき、現場の状況を明らかにするため、写真撮影を実施したことから報告する。')
title_end = doc1.add_paragraph('記')
doc1.add_paragraph()
crime_date_title = doc1.add_paragraph('１　発生日')
crime_date = doc1.add_paragraph('令和６年３月１日')



create_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
police_station.alignment = WD_ALIGN_PARAGRAPH.LEFT
police_l_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
police_l_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
crime_date.text = crime_date.text.replace('令和', '平成')
crime_place_title = doc1.add_paragraph('２　発生場所')
crime_place_base = doc1.add_paragraph('福岡市東区')
crime_place_option = crime_place_base.add_run('箱崎１丁目１番１号')
photographer_title = doc1.add_paragraph('４　撮影者')
photographer = doc1.add_paragraph('司法警察員巡査部長　東　太郎')


crime_bagage_title = doc1.add_paragraph('３　所持品')
data = [[(1), '長財布', '１個'], [(2), '黒色リュック', '１個'], [(3), 'ハンカチ', '１枚']]
table = doc1.add_table(rows=1, cols=3)
cells = table.rows[0].cells
for id, name, length in data:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = name
    row_cells[2].text = length

photographer_title.insert_paragraph_before(crime_bagage_title.text)


doc1.save('new.docx')